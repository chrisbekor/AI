import gradio as gr
from dotenv import load_dotenv
import assemblyai as aai
import librosa
import soundfile as sf
import torch
import json
import csv
import os
from datetime import datetime
from transformers import AutoTokenizer, AutoModelForSequenceClassification
import torch.nn.functional as F
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

# =========================
# CONFIGs
# =========================
load_dotenv()

aai.settings.api_key = os.getenv("ASSEMBLYAI_API_KEY")
hf_token = os.getenv("HF_TOKEN")

device = "cuda" if torch.cuda.is_available() else "cpu"

tokenizer = AutoTokenizer.from_pretrained(
    "nlptown/bert-base-multilingual-uncased-sentiment"
)
model = AutoModelForSequenceClassification.from_pretrained(
    "nlptown/bert-base-multilingual-uncased-sentiment"
)
model.to(device)
model.eval()

# =========================
# GLOBAL
# =========================
global_segments = []
global_conversation = ""

# =========================
# HELPERS
# =========================
def format_time(ms):
    s = ms / 1000
    return f"{int(s//60):02d}:{int(s%60):02d}"


def build_segments(transcript):
    speaker_map = {}   # maps A → 1, B → 2, etc.
    current_id = 1

    segments = []

    for u in transcript.utterances:
        raw_speaker = str(u.speaker)

        # assign numeric ID if not already mapped
        if raw_speaker not in speaker_map:
            speaker_map[raw_speaker] = current_id
            current_id += 1

        speaker_id = speaker_map[raw_speaker]

        segments.append({
            "speaker": speaker_id,
            "start": format_time(u.start or 0),
            "end": format_time(u.end or 0),
            "text": u.text
        })

    return segments

def analyze_text(text):
    inputs = tokenizer(text[:512], return_tensors="pt", truncation=True).to(device)

    with torch.no_grad():
        logits = model(**inputs).logits

    probs = F.softmax(logits, dim=-1)[0]
    return torch.argmax(probs).item() + 1


# =========================
# MAIN PROCESS
# =========================
def process_audio(file, speakers=0, language="auto"):
    global global_segments, global_conversation

    if file is None:
        return "❌ No audio provided", "", ""

    path = file.name if hasattr(file, "name") else file

    try:
        audio, sr = librosa.load(path, sr=None, mono=True)
        tmp = "temp.wav"
        sf.write(tmp, audio, sr)

        config = aai.TranscriptionConfig(
            speaker_labels=True,
            speakers_expected=speakers if speakers > 0 else None,
            language_code=None if language == "auto" else language
        )

        transcript = aai.Transcriber().transcribe(tmp, config)

        if transcript.error:
            return f"❌ {transcript.error}", "", ""

        global_segments = build_segments(transcript)

        # Unique speaker count
        speaker_ids = sorted(set(s["speaker"] for s in global_segments))
        speaker_count = len(speaker_ids)

        label_map = {
            1: ("🔴", "Very Negative"),
            2: ("🟠", "Negative"),
            3: ("🟡", "Neutral"),
            4: ("🟢", "Positive"),
            5: ("🟢", "Very Positive")
        }

        conversation = ""

        for i, seg in enumerate(global_segments, start=1):
            score = analyze_text(seg["text"])
            emoji, label = label_map.get(score, ("⚪", "Unknown"))

            conversation += (
                f"Speaker {seg['speaker']} | Utterance {i}\n"
                f"({seg['start']} - {seg['end']})\n"
                f"{emoji} {label}: {seg['text']}\n\n"
            )

        global_conversation = conversation

        return "✅ Done", conversation, f"Speakers: {speaker_count} | Utterances: {len(global_segments)}"

    except Exception as e:
        return f"❌ Error: {str(e)}", "", ""


# =========================
# EXPORT
# =========================
def export_file(format_type):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    if format_type == "TXT":
        path = f"conversation_{timestamp}.txt"
        with open(path, "w", encoding="utf-8") as f:
            f.write(global_conversation)

    elif format_type == "JSON":
        path = f"conversation_{timestamp}.json"
        with open(path, "w", encoding="utf-8") as f:
            json.dump(global_segments, f, indent=4)

    elif format_type == "CSV":
        path = f"conversation_{timestamp}.csv"
        with open(path, "w", newline="", encoding="utf-8") as f:
            writer = csv.DictWriter(f, fieldnames=["speaker", "start", "end", "text"])
            writer.writeheader()
            writer.writerows(global_segments)

    elif format_type == "WORD":
        path = f"conversation_{timestamp}.docx"
        doc = Document()
        doc.add_heading("Conversation Transcript", 0)
        doc.add_paragraph(global_conversation)
        doc.save(path)

    elif format_type == "PDF":
        path = f"conversation_{timestamp}.pdf"
        doc = SimpleDocTemplate(path)
        styles = getSampleStyleSheet()
        content = [Paragraph(global_conversation.replace("\n", "<br/>"), styles["Normal"])]
        doc.build(content)

    return path


# =========================
# UI
# =========================
with gr.Blocks(title="AI Conversation Sentiment System") as app:

    gr.Markdown("# 🎙 AI Conversation Sentiment Analyzer")

    with gr.Group():
        gr.Markdown("### 🎙 Input Audio")
        audio = gr.Audio(sources=["upload", "microphone"], type="filepath")

    with gr.Group():
        gr.Markdown("### ⚙ Settings")
        with gr.Row():
            speakers = gr.Number(value=0, label="Speakers (0 = auto)")
            language = gr.Dropdown(["auto", "en", "fr", "es", "de"], value="auto")

    analyze_btn = gr.Button("🚀 Analyze")

    with gr.Group():
        gr.Markdown("### 💬 Conversation Output")
        status = gr.Textbox(label="Status")
        conversation_box = gr.Textbox(lines=18, label="Conversation + Sentiment")
        info = gr.Textbox(label="Info")

    with gr.Group():
        gr.Markdown("### 📁 Export")
        with gr.Row():
            export_format = gr.Dropdown(
                ["TXT", "JSON", "CSV", "WORD", "PDF"],
                value="TXT",
                label="Select Format"
            )
            export_btn = gr.Button("⬇ Export")
        download = gr.File()

    analyze_btn.click(
        process_audio,
        inputs=[audio, speakers, language],
        outputs=[status, conversation_box, info]
    )

    export_btn.click(
        export_file,
        inputs=[export_format],
        outputs=[download]
    )


if __name__ == "__main__":
    app.launch(theme=gr.themes.Soft())